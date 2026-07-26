"""
ado_lib — shared library for the ado-sync CLI tools.

Consolidated from: ado_client, field_formatter, context_loader,
                   doc_ingester, ai_interpreter, project_builder,
                   and the shared execution helpers from ai_ado_creator.

Entry points (ai_ado_creator.py, ado_repo_sync.py) import from here.
All connection config comes from environment variables — no hardcoded defaults.
Required env vars: ADO_ORG, ADO_PROJECT, ADO_PAT, ANTHROPIC_API_KEY
Optional: ADO_API_VERSION (7.1), ADO_PROJECT_GUID, ADO_CONTEXT_CACHE_PATH, ADO_PLAN_DIR
"""
from __future__ import annotations

import html
import io
import json
import os
import re
import subprocess
import sys
import time
import urllib.parse

# ── Auto-install third-party deps ─────────────────────────────────────────────
for _pkg in ("anthropic", "requests"):
    try:
        __import__(_pkg)
    except ModuleNotFoundError:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", _pkg, "-q", "--break-system-packages"],
            stdout=subprocess.DEVNULL,
        )

import anthropic  # type: ignore[import]
import requests   # type: ignore[import]


# ══════════════════════════════════════════════════════════════════════════════
# ADO Client
# ══════════════════════════════════════════════════════════════════════════════

ORG          = os.getenv("ADO_ORG", "")
PROJECT      = os.getenv("ADO_PROJECT", "")
PAT          = os.getenv("ADO_PAT", "")
API_VERSION  = os.getenv("ADO_API_VERSION", "7.1")
PROJECT_GUID = os.getenv("ADO_PROJECT_GUID", "")


class ADOClient:
    """Thin wrapper around the Azure DevOps REST API."""

    def __init__(self, org=ORG, project=PROJECT, pat=PAT):
        if not pat:
            raise ValueError(
                "ADO_PAT is not set. "
                "Set ADO_PAT in your environment or .env before running ado-sync tools."
            )
        self.org         = org
        self.project     = project
        self.pat         = pat
        self.enc_project = urllib.parse.quote(project, safe="")
        self.base_url    = f"https://dev.azure.com/{org}/{self.enc_project}"
        self.auth        = ("", pat)

    def _get(self, url, timeout=30, **kwargs):
        r = requests.get(url, auth=self.auth, timeout=timeout, **kwargs)
        r.raise_for_status()
        return r.json()

    def _post_json(self, url, data):
        r = requests.post(
            url, auth=self.auth, json=data,
            headers={"Content-Type": "application/json"},
            timeout=30,
        )
        r.raise_for_status()
        return r.json()

    def _post_patch(self, url, ops):
        r = requests.post(
            url, auth=self.auth, json=ops,
            headers={"Content-Type": "application/json-patch+json"},
            timeout=30,
        )
        r.raise_for_status()
        return r.json()

    def _patch(self, url, ops):
        r = requests.patch(
            url, auth=self.auth, json=ops,
            headers={"Content-Type": "application/json-patch+json"},
            timeout=30,
        )
        r.raise_for_status()
        return r.json()

    # ── Reads ──────────────────────────────────────────────────────────────────

    def get_work_item(self, item_id, fields=None):
        url    = f"{self.base_url}/_apis/wit/workitems/{item_id}"
        params = {"api-version": API_VERSION}
        if fields:
            params["fields"] = ",".join(fields)
        return self._get(url, params=params)

    def get_work_items_batch(self, ids, fields=None):
        if not ids:
            return []
        url    = f"{self.base_url}/_apis/wit/workitems"
        params = {"ids": ",".join(str(i) for i in ids), "api-version": API_VERSION}
        if fields:
            params["fields"] = ",".join(fields)
        return self._get(url, params=params).get("value", [])

    def run_wiql(self, query):
        url = f"{self.base_url}/_apis/wit/wiql?api-version={API_VERSION}"
        return self._post_json(url, {"query": query})

    def find_work_item(self, wit_type, title, parent_id=None):
        safe_title = title.replace("'", "''")
        result = self.run_wiql(f"""
            SELECT [System.Id] FROM WorkItems
            WHERE [System.TeamProject] = '{self.project}'
              AND [System.WorkItemType] = '{wit_type}'
              AND [System.Title] = '{safe_title}'
              AND [System.State] <> 'Removed'
        """)
        items = result.get("workItems", [])
        if not items:
            return None
        if parent_id and len(items) > 1:
            for item in items:
                wi = self.get_work_item(item["id"], fields=["System.Parent"])
                if wi["fields"].get("System.Parent") == parent_id:
                    return item["id"]
        return items[0]["id"]

    def get_project_hierarchy(self):
        result = self.run_wiql(f"""
            SELECT [System.Id] FROM WorkItems
            WHERE [System.TeamProject] = '{self.project}'
              AND [System.WorkItemType] IN (
                  'Epic','Feature','Product Backlog Item','Automation Request'
              )
              AND [System.State] <> 'Removed'
            ORDER BY [System.WorkItemType], [System.Id]
        """)
        ids    = [i["id"] for i in result.get("workItems", [])]
        fields = [
            "System.Id", "System.Title", "System.WorkItemType",
            "System.State", "System.Parent", "System.Tags", "System.AreaPath",
        ]
        items = []
        for i in range(0, len(ids), 200):
            items.extend(self.get_work_items_batch(ids[i:i + 200], fields=fields))
        return items

    def get_child_features(self, epic_id: int) -> list[dict]:
        result = self.run_wiql(f"""
            SELECT [System.Id] FROM WorkItems
            WHERE [System.TeamProject] = '{self.project}'
              AND [System.WorkItemType] = 'Feature'
              AND [System.Parent] = {epic_id}
              AND [System.State] <> 'Removed'
        """)
        ids = [i["id"] for i in result.get("workItems", [])]
        if not ids:
            return []
        return self.get_work_items_batch(ids, fields=["System.Id", "System.Title", "System.State"])

    # ── Writes ─────────────────────────────────────────────────────────────────

    def create_work_item(self, wit_type, ops):
        enc_type = urllib.parse.quote(wit_type, safe="")
        url      = f"{self.base_url}/_apis/wit/workitems/${enc_type}?api-version={API_VERSION}"
        return self._post_patch(url, ops)

    def update_work_item(self, item_id, ops):
        url = f"{self.base_url}/_apis/wit/workitems/{item_id}?api-version={API_VERSION}"
        return self._patch(url, ops)

    # ── Patch helpers ──────────────────────────────────────────────────────────

    @staticmethod
    def field_op(field, value):
        return {"op": "add", "path": f"/fields/{field}", "value": value}

    def parent_link_op(self, parent_id, comment=""):
        return {
            "op": "add",
            "path": "/relations/-",
            "value": {
                "rel": "System.LinkTypes.Hierarchy-Reverse",
                "url": (f"https://dev.azure.com/{self.org}/{PROJECT_GUID}"
                        f"/_apis/wit/workItems/{parent_id}"),
                "attributes": {"comment": comment},
            },
        }


# ══════════════════════════════════════════════════════════════════════════════
# Field Formatter
# ══════════════════════════════════════════════════════════════════════════════

FIELD_FORMAT = {
    "System.Title":                             "plain",
    "System.Description":                       "html",
    "System.AcceptanceCriteria":                "html",
    "Microsoft.VSTS.Common.AcceptanceCriteria": "html",
    "Microsoft.VSTS.TCM.ReproSteps":            "html",
    "System.Tags":                              "plain",
    "System.State":                             "plain",
    "System.AreaPath":                          "plain",
    "System.AssignedTo":                        "plain",
    "Microsoft.VSTS.Common.Priority":           "integer",
    "Microsoft.VSTS.Common.ValueArea":          "plain",
}


def md_to_html(text: str) -> str:
    """Convert markdown to HTML for ADO rich-text fields. Passes through existing HTML."""
    if not text or not text.strip():
        return ""
    if re.search(r"<(p|h[1-6]|ul|ol|li|br|div|strong|em|code|pre)\b", text, re.IGNORECASE):
        return text

    def inline(s: str) -> str:
        s = html.escape(s, quote=False)
        s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
        s = re.sub(r"\*(.+?)\*",     r"<em>\1</em>",         s)
        s = re.sub(r"`(.+?)`",       r"<code>\1</code>",     s)
        return s

    out   = []
    in_ul = in_ol = False

    def close_lists():
        nonlocal in_ul, in_ol
        if in_ul:
            out.append("</ul>")
            in_ul = False
        if in_ol:
            out.append("</ol>")
            in_ol = False

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            close_lists()
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            close_lists()
            out.append(f"<h{len(m.group(1))}>{inline(m.group(2))}</h{len(m.group(1))}>")
            continue
        m = re.match(r"^[-*•]\s+(.*)", line)
        if m:
            if not in_ul:
                close_lists(); out.append("<ul>"); in_ul = True
            out.append(f"<li>{inline(m.group(1))}</li>")
            continue
        m = re.match(r"^\d+[.)]\s+(.*)", line)
        if m:
            if not in_ol:
                close_lists(); out.append("<ol>"); in_ol = True
            out.append(f"<li>{inline(m.group(1))}</li>")
            continue
        close_lists()
        out.append(f"<p>{inline(line)}</p>")

    close_lists()
    return "\n".join(out)


def format_field(field_ref: str, value) -> object:
    """Return value formatted for the given ADO field reference."""
    fmt = FIELD_FORMAT.get(field_ref, "plain")
    if fmt == "html" and isinstance(value, str):
        return md_to_html(value)
    if fmt == "integer":
        return int(value)
    return value


# ══════════════════════════════════════════════════════════════════════════════
# Context Loader
# ══════════════════════════════════════════════════════════════════════════════

CACHE_PATH = os.getenv(
    "ADO_CONTEXT_CACHE_PATH",
    os.path.join(os.path.dirname(__file__), ".ado_context_cache.json"),
)
CACHE_TTL = 600  # 10 minutes


def load_context(ado_client, force_refresh=False):
    """Return list of work item dicts, refreshing cache when stale."""
    if not force_refresh and os.path.exists(CACHE_PATH):
        with open(CACHE_PATH, encoding="utf-8") as f:
            cached = json.load(f)
        if time.time() - cached.get("timestamp", 0) < CACHE_TTL:
            return cached["items"]

    print("  Refreshing ADO context from live data...", flush=True)
    items = ado_client.get_project_hierarchy()
    with open(CACHE_PATH, "w", encoding="utf-8") as f:
        json.dump({"timestamp": time.time(), "items": items}, f)
    return items


def build_context_summary(items: list) -> str:
    """Build a compact text block for the AI system prompt from the live work item hierarchy."""
    epics    = [i for i in items if i["fields"]["System.WorkItemType"] == "Epic"]
    features = [i for i in items if i["fields"]["System.WorkItemType"] == "Feature"]
    pbis     = [i for i in items if i["fields"]["System.WorkItemType"] == "Product Backlog Item"]
    reqs     = [i for i in items if i["fields"]["System.WorkItemType"] == "Automation Request"]

    epic_map    = {e["id"]: e["fields"]["System.Title"] for e in epics}
    feature_map = {f["id"]: f["fields"]["System.Title"] for f in features}

    lines = ["## Live ADO Work Item Hierarchy\n"]
    lines.append("### Epics")
    for e in epics:
        f = e["fields"]
        lines.append(f"- [#{f['System.Id']}] {f['System.Title']}  ({f.get('System.State', '')})")

    lines.append("\n### Features")
    for feat in features:
        f = feat["fields"]
        lines.append(
            f"- [#{f['System.Id']}] {f['System.Title']}  "
            f"({f.get('System.State', '')})  → Epic: {epic_map.get(f.get('System.Parent'), '—')}"
        )

    lines.append("\n### Product Backlog Items")
    for pbi in pbis:
        f = pbi["fields"]
        lines.append(
            f"- [#{f['System.Id']}] {f['System.Title']}  "
            f"({f.get('System.State', '')})  → Feature: {feature_map.get(f.get('System.Parent'), '—')}"
        )

    if reqs:
        lines.append("\n### Automation Requests")
        for ar in reqs:
            f = ar["fields"]
            lines.append(f"- [#{f['System.Id']}] {f['System.Title']}  ({f.get('System.State', '')})")

    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# Doc Ingester
# ══════════════════════════════════════════════════════════════════════════════

SUPPORTED_DOC_TYPES = {".pdf", ".docx", ".txt", ".md", ".xlsx"}
MAX_CHARS_PER_FILE  = 40_000
MAX_TOTAL_CHARS     = 150_000


def _read_pdf(data: bytes) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as exc:
        return f"[PDF extract error: {exc}]"


def _read_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc   = Document(io.BytesIO(data))
        parts = []

        def _cell_text(cell) -> str:
            return " | ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())

        def _table_text(table) -> str:
            rows = []
            for row in table.rows:
                row_text = "  ".join(c for c in [_cell_text(c) for c in row.cells] if c)
                if row_text:
                    rows.append(row_text)
            return "\n".join(rows)

        for block in doc.element.body:
            tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
            if tag == "p":
                text = "".join(node.text or "" for node in block.iter() if node.tag.endswith("}t"))
                if text.strip():
                    parts.append(text.strip())
            elif tag == "tbl":
                from docx.table import Table
                t = _table_text(Table(block, doc))
                if t:
                    parts.append(t)

        for section in doc.sections:
            for hf in [section.header, section.footer]:
                if hf:
                    for p in hf.paragraphs:
                        if p.text.strip():
                            parts.append(p.text.strip())
        return "\n".join(parts)
    except Exception as exc:
        return f"[DOCX extract error: {exc}]"


def _read_xlsx(data: bytes) -> str:
    try:
        import openpyxl
        wb    = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append("  " + " | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        return f"[XLSX extract error: {exc}]"


def _extract_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    with open(file_path, "rb") as f:
        data = f.read()
    if ext == ".pdf":
        return _read_pdf(data)
    if ext == ".docx":
        return _read_docx(data)
    if ext == ".xlsx":
        return _read_xlsx(data)
    return data.decode("utf-8", errors="replace")


def ingest(path: str) -> tuple[str, list[str]]:
    """
    Extract text from a single file or all supported files in a folder.
    Returns (combined_text, list_of_file_names).
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"Path not found: {path}")
    if os.path.isdir(path):
        entries = sorted(
            f for f in os.listdir(path)
            if os.path.splitext(f)[1].lower() in SUPPORTED_DOC_TYPES and not f.startswith(".")
        )
        if not entries:
            raise ValueError(
                f"No supported files in folder. Supported: {', '.join(sorted(SUPPORTED_DOC_TYPES))}"
            )
        sections, processed, total = [], [], 0
        for name in entries:
            text = _extract_file(os.path.join(path, name))[:MAX_CHARS_PER_FILE]
            sections.append(f"=== FILE: {name} ===\n{text}\n")
            processed.append(name)
            total += len(text)
            if total >= MAX_TOTAL_CHARS:
                sections.append("[Remaining files omitted — character limit reached]")
                break
        return "\n".join(sections), processed

    ext = os.path.splitext(path)[1].lower()
    if ext not in SUPPORTED_DOC_TYPES:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_DOC_TYPES))}"
        )
    name = os.path.basename(path)
    return f"=== FILE: {name} ===\n{_extract_file(path)[:MAX_CHARS_PER_FILE]}\n", [name]


# ══════════════════════════════════════════════════════════════════════════════
# AI Interpreter  (plain-English → ADO operations)
# ══════════════════════════════════════════════════════════════════════════════

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

_INTERPRETER_SYSTEM_PROMPT = """\
You are an Azure DevOps (ADO) work item assistant.

## Your job
Parse a plain-English progress update or request into a precise list of ADO work item operations.
Always call the `submit_operations` tool — never return raw JSON or prose.

## Work Item Model
| Level       | Type                   | Common States                                    |
|-------------|------------------------|--------------------------------------------------|
| Strategy    | Epic                   | To Do · Ready · Work in Progress · Done          |
| Capability  | Feature                | Draft · Active · Closed                          |
| Intake      | Automation Request     | New · Ready for Analysis · In Analysis · Closed  |
| Delivery    | Product Backlog Item   | New · Active · Closed                            |
| Task        | Task                   | To Do · In Progress · Done                       |
| Defect      | Bug                    | New · Active · Resolved · Closed                 |
| Risk/Block  | Issue                  | New · Active · Closed                            |
| Reuse       | Automation Asset       | New · Active · Closed                            |
| Operations  | Support Item           | New · Active · Closed                            |

## SDLC Phases
Intake → Analysis → Build → Test → Release → Support

## Inference Rules
1. "finished" / "completed" / "done" / "closed out"
   → Task state: Done  |  PBI state: Closed
2. "started" / "working on" / "in progress" / "picking up"
   → Task state: In Progress  |  PBI state: Active
3. "blocked" / "waiting on sign-off" / "pending approval"
   → Tag the relevant item with "blocked" AND consider creating an Issue
4. "need to" / "plan to" / "next step" / "todo"
   → Task state: To Do  |  PBI state: New
5. Epics are referenced by code (e.g. PROJ-001) — match to the live hierarchy
6. Prefer updating existing items over creating new ones when the title is an obvious match
7. When creating a Task, always supply parent_id (the PBI) or parent_title + parent_type
8. When creating a PBI, always supply parent_id (the Feature) or parent_title + parent_type
9. If the user's input is ambiguous, create the most specific item the context supports
   and explain the assumption in `reasoning`

## Field Format Rules (informational — formatter handles conversion)
- System.Title → plain text
- System.Description → HTML (you may write markdown; it will be converted)
- System.AcceptanceCriteria → HTML
- System.Tags → plain text, semicolon-separated
"""

_SUBMIT_TOOL = {
    "name": "submit_operations",
    "description": "Submit ADO work item operations to execute against the live project.",
    "input_schema": {
        "type": "object",
        "required": ["operations", "summary"],
        "properties": {
            "operations": {
                "type": "array",
                "items": {
                    "type": "object",
                    "required": ["action", "work_item_type", "title"],
                    "properties": {
                        "action":          {"type": "string", "enum": ["create", "update", "find"]},
                        "work_item_type":  {
                            "type": "string",
                            "enum": [
                                "Epic", "Feature", "Automation Request",
                                "Product Backlog Item", "Task",
                                "Bug", "Issue", "Automation Asset", "Support Item",
                            ],
                        },
                        "work_item_id":    {"type": "integer"},
                        "title":           {"type": "string"},
                        "description":     {"type": "string"},
                        "acceptance_criteria": {"type": "string"},
                        "state":           {"type": "string"},
                        "parent_id":       {"type": "integer"},
                        "parent_title":    {"type": "string"},
                        "parent_type":     {"type": "string"},
                        "tags":            {"type": "string"},
                        "area_path":       {"type": "string"},
                        "reasoning":       {"type": "string"},
                    },
                },
            },
            "summary": {"type": "string"},
        },
    },
}


def interpret(user_input: str, context_summary: str) -> dict:
    """Send plain-English input + live ADO context to Claude; return submit_operations dict."""
    client   = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    response = client.messages.create(
        model="claude-haiku-4-5",
        max_tokens=2048,
        system=[
            {"type": "text", "text": _INTERPRETER_SYSTEM_PROMPT, "cache_control": {"type": "ephemeral"}},
            {"type": "text", "text": context_summary,            "cache_control": {"type": "ephemeral"}},
        ],
        tools=[_SUBMIT_TOOL],
        tool_choice={"type": "tool", "name": "submit_operations"},
        messages=[{"role": "user", "content": user_input}],
    )
    for block in response.content:
        if block.type == "tool_use" and block.name == "submit_operations":
            return block.input
    raise ValueError("Claude did not return a submit_operations call")


# ══════════════════════════════════════════════════════════════════════════════
# Project Builder  (documents → full ADO hierarchy)
# ══════════════════════════════════════════════════════════════════════════════

_BUILDER_SYSTEM_PROMPT = """\
You are an Azure DevOps project architect.

## Your job
Read the provided project requirements document(s) and generate a complete, ready-to-execute
ADO work item hierarchy. Always call the `submit_operations` tool — never return raw JSON or prose.

## Work Item Model
Every project follows this standard structure:

Epic (one per project)
├── Feature: Enhancements & New Capabilities   ← all build/dev PBIs go here
├── Feature: Defects & Production Issues
├── Feature: Tech Debt & Refactoring
└── Feature: Runbooks, Monitoring & Operations

## Feature Naming Rules
The user message will tell you whether an "Initial Build" feature already exists under the target Epic.

- **If no "Initial Build" feature exists** → create exactly one Feature titled `Initial Build`.
  All development PBIs go under this Feature.

- **If "Initial Build" already exists** → do NOT create an "Initial Build" feature.
  Instead, read the document content and determine whether this request is primarily a defect fix
  or an enhancement/new capability, then name the Feature accordingly:
  - Defect fix → `Defect – <concise summary of the issue>` (≤ 8 words after the dash)
  - Enhancement / new capability → `Enhancement – <concise title>` (≤ 8 words after the dash)

Always create exactly **one** Feature (not multiple). All PBIs and tasks nest under it.

## Epic Naming Convention
Format: `<CODE> – <Short Business Name>`
Examples: `PROJ-001 – Customer Onboarding Automation`, `PROJ-002 – Invoice Processing`

If the documents don't specify a code, use `PROJ` as the prefix.

## PBI Guidelines
- Each PBI should represent a single meaningful deliverable (e.g., "Discovery & Solution Design",
  "Automate [Core Process]", "Integration Testing & UAT", "Deployment & Runbook")
- Always start with a "Discovery & Solution Design" PBI under Enhancements
- Include concrete acceptance criteria for each PBI
- Initial state: New

## Task Guidelines
- 4–8 tasks per PBI, representing concrete work steps
- Cover: design, implementation, peer review, testing
- Initial state: To Do

## State Reference
- Epic → "To Do"
- Feature → "Draft"
- Product Backlog Item → "New"
- Task → "To Do"

## Field Format Rules
- System.Title → plain text, concise
- System.Description → markdown (auto-converted to HTML)
- System.AcceptanceCriteria → markdown bullet list of testable criteria
- System.Tags → semicolon-separated

## Operation Order
Submit in this order so parents exist before children:
1. Epic (create)
2. Features (create, with parent_title = Epic title)
3. PBIs (create, with parent_title = Feature title)
4. Tasks (create, with parent_title = PBI title)

Always set parent_title AND parent_type on every child item.
"""

_SCAFFOLD_PROMPT = """\
You are an Azure DevOps project architect.

## Your job
Read the provided project request or document(s) and generate a MINIMAL ADO scaffold
to kick off discovery work. Always call the `submit_operations` tool.

## What to create (nothing more, nothing less)
1. One **Epic** — use the exact name provided in the user message as `EPIC_NAME`
2. One **Feature** — title exactly: `Discovery`, child of the Epic
3. One **Product Backlog Item** — title: `Discovery & Solution Design`, child of the Feature
   - Write a description summarising the project request in 2–3 sentences
   - Write acceptance criteria as a bullet list covering: process documented, systems identified,
     data flows mapped, solution design reviewed and signed off
4. **5–7 Tasks** under that PBI covering:
   - Document current process and pain points
   - Identify systems, applications, and data sources involved
   - Map data flows, triggers, and edge cases
   - Define automation scope and out-of-scope boundaries
   - Draft solution design document
   - Schedule and conduct design review / sign-off
   - (Optional) Identify reusable Automation Assets

## Rules
- State for Epic: To Do  |  Feature: Draft  |  PBI: New  |  Tasks: To Do
- Every child must have parent_title AND parent_type set
- Emit operations in order: Epic → Feature → PBI → Tasks
- Do NOT create additional Features, PBIs, or Tasks beyond this structure
- Keep titles concise (plain text, no markdown)
- Description and acceptance_criteria may use markdown (auto-converted to HTML)
"""

_BUILDER_SUBMIT_TOOL = {
    "name": "submit_operations",
    "description": "Submit the full ADO work item hierarchy to create from the project documents.",
    "input_schema": {
        "type": "object",
        "required": ["operations", "summary"],
        "properties": {
            "operations": {
                "type": "array",
                "items": {
                    "type": "object",
                    "required": ["action", "work_item_type", "title"],
                    "properties": {
                        "action":              {"type": "string", "enum": ["create", "update", "find"]},
                        "work_item_type":      {"type": "string", "enum": [
                            "Epic", "Feature", "Automation Request",
                            "Product Backlog Item", "Task",
                            "Bug", "Issue", "Automation Asset", "Support Item",
                        ]},
                        "work_item_id":        {"type": "integer"},
                        "title":               {"type": "string"},
                        "description":         {"type": "string"},
                        "acceptance_criteria": {"type": "string"},
                        "state":               {"type": "string"},
                        "parent_id":           {"type": "integer"},
                        "parent_title":        {"type": "string"},
                        "parent_type":         {"type": "string"},
                        "tags":                {"type": "string"},
                        "area_path":           {"type": "string"},
                        "reasoning":           {"type": "string"},
                    },
                },
            },
            "summary": {"type": "string"},
        },
    },
}


def _call_claude(system_prompt: str, user_message: str, max_tokens: int = 4096) -> dict:
    client   = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=max_tokens,
        system=[{"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}}],
        tools=[_BUILDER_SUBMIT_TOOL],
        tool_choice={"type": "tool", "name": "submit_operations"},
        messages=[{"role": "user", "content": user_message}],
    )
    print(f"  [AI] stop_reason={response.stop_reason}  "
          f"tokens: {response.usage.input_tokens} in / {response.usage.output_tokens} out")
    if response.stop_reason == "max_tokens":
        print("  [AI] WARNING: output truncated — increase max_tokens or reduce document size.")
    for block in response.content:
        if block.type == "tool_use" and block.name == "submit_operations":
            ops = block.input.get("operations", [])
            if not ops:
                print(f"  [AI] Tool called but returned 0 operations. "
                      f"Summary: {block.input.get('summary', '(none)')}")
            return block.input
    raise ValueError("Claude did not return a submit_operations tool call")


def build_scaffold(text: str, epic_name: str, source_label: str = "project request") -> dict:
    """Generate a minimal discovery scaffold (Epic + Discovery Feature + 1 PBI + tasks)."""
    return _call_claude(
        _SCAFFOLD_PROMPT,
        f"EPIC_NAME: {epic_name}\n\nProject request / document ({source_label}):\n\n--- CONTENT ---\n{text}",
        max_tokens=4096,
    )


def build_from_documents(
    doc_text: str,
    file_names: list[str],
    epic_name: str = "",
    has_initial_build: bool = False,
    existing_features: list[str] | None = None,
) -> dict:
    """Send extracted document content to Claude; return a full ADO hierarchy."""
    feature_context = ""
    if epic_name:
        feature_context += f"TARGET EPIC: {epic_name}\n"
    if has_initial_build:
        feature_context += (
            "EXISTING FEATURES: An 'Initial Build' feature already exists under this Epic.\n"
            "Do NOT create an 'Initial Build' feature. Instead, name the new Feature either\n"
            "'Defect – <summary>' or 'Enhancement – <title>' based on the document content.\n"
        )
        if existing_features:
            feature_context += f"All existing features: {', '.join(existing_features)}\n"
    else:
        feature_context += (
            "EXISTING FEATURES: No 'Initial Build' feature exists yet.\n"
            "Create the Feature with the title 'Initial Build'.\n"
        )
    user_message = (
        f"{feature_context}\n"
        f"I have uploaded the following project document(s): {', '.join(file_names)}\n\n"
        f"Please analyse them and generate a complete ADO work item hierarchy.\n\n"
        f"--- DOCUMENT CONTENT ---\n{doc_text}"
    )
    return _call_claude(_BUILDER_SYSTEM_PROMPT, user_message, max_tokens=16000)


# ══════════════════════════════════════════════════════════════════════════════
# Operation Execution  (shared between ai_ado_creator and ado_repo_sync)
# ══════════════════════════════════════════════════════════════════════════════

ADO_WEB  = f"https://dev.azure.com/{ORG}/{PROJECT.replace(' ', '%20')}/_workitems/edit"
PLAN_DIR = os.getenv("ADO_PLAN_DIR", "./_engineer/ado-sync/plans")
PLAN_FILE = os.path.join(PLAN_DIR, "ado_last_plan.json")

_DEFAULT_STATE: dict[str, str] = {
    "Epic":                 "To Do",
    "Feature":              "Draft",
    "Product Backlog Item": "New",
    "Task":                 "To Do",
    "Bug":                  "New",
    "Issue":                "New",
    "Automation Request":   "New",
    "Automation Asset":     "New",
    "Support Item":         "New",
}


def _coerce_state(wit_type: str, state: str) -> str:
    valid = {
        "Epic":    {"To Do", "Ready", "Work in Progress", "Done", "Removed"},
        "Feature": {"Draft", "Active", "Closed", "Removed"},
    }
    allowed = valid.get(wit_type)
    if allowed and state not in allowed:
        return _DEFAULT_STATE.get(wit_type, state)
    return state


def _resolve_parent_id(ado: ADOClient, op: dict) -> int | None:
    if op.get("parent_id"):
        return op["parent_id"]
    if op.get("parent_title") and op.get("parent_type"):
        return ado.find_work_item(op["parent_type"], op["parent_title"])
    return None


def _build_create_patches(ado: ADOClient, op: dict) -> list:
    patches = [ado.field_op("System.Title", op["title"])]
    if op.get("description"):
        patches.append(ado.field_op("System.Description",
                                    format_field("System.Description", op["description"])))
    if op.get("acceptance_criteria"):
        patches.append(ado.field_op("Microsoft.VSTS.Common.AcceptanceCriteria",
                                    format_field("Microsoft.VSTS.Common.AcceptanceCriteria",
                                                 op["acceptance_criteria"])))
    if op.get("state"):
        patches.append(ado.field_op("System.State",
                                    _coerce_state(op.get("work_item_type", ""), op["state"])))
    if op.get("tags"):
        patches.append(ado.field_op("System.Tags", op["tags"]))
    if op.get("area_path"):
        patches.append(ado.field_op("System.AreaPath", op["area_path"]))
    parent_id = _resolve_parent_id(ado, op)
    if parent_id:
        patches.append(ado.parent_link_op(parent_id))
    return patches


def _build_update_patches(op: dict) -> list:
    patches   = []
    field_map = {
        "title":               "System.Title",
        "description":         "System.Description",
        "acceptance_criteria": "Microsoft.VSTS.Common.AcceptanceCriteria",
        "state":               "System.State",
        "tags":                "System.Tags",
        "area_path":           "System.AreaPath",
    }
    for key, field_ref in field_map.items():
        if op.get(key):
            patches.append({
                "op": "add", "path": f"/fields/{field_ref}",
                "value": format_field(field_ref, op[key]),
            })
    return patches


def execute_operation(ado: ADOClient, op: dict, dry_run: bool) -> None:
    """Execute a single ADO operation (find / create / update)."""
    action = op.get("action", "").lower()
    wtype  = op.get("work_item_type", "?")
    title  = op.get("title", "?")

    if action == "find":
        found = ado.find_work_item(wtype, title)
        if found:
            print(f"     Found  #{found}: [{wtype}] {title}")
        else:
            print(f"     Not found: [{wtype}] {title!r}")
        return

    if action == "create":
        patches = _build_create_patches(ado, op)
        if dry_run:
            print(f"     [DRY RUN] CREATE {wtype}: \"{title}\"")
            print(f"              patches: {json.dumps(patches, indent=14)}")
            return
        result  = ado.create_work_item(wtype, patches)
        item_id = result["id"]
        print(f"     Created #{item_id}  {ADO_WEB}/{item_id}")
        return

    if action == "update":
        item_id = op.get("work_item_id") or ado.find_work_item(wtype, title)
        if not item_id:
            print(f"     Could not find [{wtype}] \"{title}\" — skipped.")
            return
        patches = _build_update_patches(op)
        if not patches:
            print(f"     Nothing to update on #{item_id} — skipped.")
            return
        if dry_run:
            print(f"     [DRY RUN] UPDATE #{item_id} {wtype}: \"{title}\"")
            print(f"              patches: {json.dumps(patches, indent=14)}")
            return
        ado.update_work_item(item_id, patches)
        print(f"     Updated #{item_id}  {ADO_WEB}/{item_id}")
        return

    print(f"     Unknown action '{action}' — skipped.")


def execute_operation_returning_id(ado: ADOClient, op: dict, dry_run: bool) -> int | None:
    """Like execute_operation but returns the new ADO item ID on create."""
    if op.get("action", "").lower() == "create":
        patches = _build_create_patches(ado, op)
        if dry_run:
            print(f"     [DRY RUN] CREATE {op.get('work_item_type', '?')}: \"{op.get('title', '?')}\"")
            return None
        result  = ado.create_work_item(op["work_item_type"], patches)
        item_id = result["id"]
        print(f"     Created #{item_id}  [{op['work_item_type']}] {op['title']}  {ADO_WEB}/{item_id}")
        return item_id
    execute_operation(ado, op, dry_run)
    return None


def print_proposed(ops: list) -> None:
    """Print a numbered summary of proposed ADO operations."""
    for i, op in enumerate(ops, 1):
        action    = op.get("action", "?").upper()
        wtype     = op.get("work_item_type", "?")
        title     = op.get("title", "?")
        state     = f"  → {op['state']}"       if op.get("state")        else ""
        parent    = (
            f"  under: {op['parent_title']}"   if op.get("parent_title") else
            f"  under: #{op['parent_id']}"     if op.get("parent_id")    else ""
        )
        item_ref  = f" #{op['work_item_id']}"  if op.get("work_item_id") else ""
        reasoning = f"\n     ↳ {op['reasoning']}" if op.get("reasoning") else ""
        print(f"  {i:>2}. {action}{item_ref}  [{wtype}] \"{title}\"{state}{parent}{reasoning}")
